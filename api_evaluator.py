import os 
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document
import json
import time

load_dotenv()
API_KEY = os.getenv("GROQ_API_KEY")

if not API_KEY:
    raise ValueError("Api key not found!!")

client = Groq(api_key = API_KEY)
model_name = 'llama-3.3-70b-versatile'

### Schemas (blueprints)
class jobDescription(BaseModel):
    role:'str'
    required_skills : list[str]
    preffered_skills : list[str]
    minimum_experience : float | None
    educational_requirement: list[str]
    responsibilities: list[str]

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str| None
    email: str | None
    phone: str | None
    total_experience_year: float | None
    skills : list[str] = []
    experiences: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

class Match_result(BaseModel):
    score: float 
    details: dict

### Reading the resumes 
def read_resume(file_obj, filename: str)-> str| None:
    ## function have the file object, sees if it is pdf or word and read it accordingly and return the contains
    ext = filename.lower()
    if ext.endswith(".pdf"):
        return read_pdf(file_obj)
    elif ext.endswith(".docx"):
        return read_docx(file_obj)
    else:
        return None

def read_pdf(file_obj) -> str:
    ## read the resume of pdf format and return the content in string format
    text = ""
    content = PdfReader(file_obj)
    for page in content.pages:
        if page:
            extracted = page.extract_text()
            if extracted:
                text += extracted + '\n'
    return text

def read_docx(file_obj) -> str:
    ## read the resume of docx format and return the content in string format
    text = ""
    content = Document(file_obj)
    for paragraph in content.paragraphs:
        if paragraph:
            text += paragraph.text + '\n'
    
    for table in content.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + '\n'
    return text

### Functions to interact with LLM 

def parse_job_description(job_text : str)-> jobDescription:
    ## Extracts structured json format file from specified raw job description
    system_prompt = f'''You are an expert HR assistant.
    Your job is to analyze job descriptions and extract structured information from them.
    Return ONLY valid JSON matching this schema:
    {jobDescription.model_json_schema()}
    
    IMPORTANT:
    Do NOT return the schema itself. Do NOT return fields like "properties", "title" or "type".
    Fill the schema with actual information extracted from the job description.
    If minimum experience is not mentioned, return null.
    If information for a list is missing, return an empty list.
    Do not invent information.'''

    messages = [
        { "role": 'system', 'content': system_prompt},
        {'role': "user", 'content': f"analyze the following job descriptions and do the steps you are told:\n\n{job_text}"}
    ]

    response = client.chat.completions.create(model = model_name, messages = messages, response_format = {"type": "json_object"})
    raw_json = response.choices[0].message.content
    job_data = json.loads(raw_json)
    return  jobDescription(**job_data)


def parse_resume(resume_text: str) -> Resume:
    """Extracts structured JSON data from a raw Resume."""
    system_prompt = f"""
    You are an expert resume parser.
    Extract information from the resume based on its meaning, not only based on exact section headings.
    Return ONLY valid JSON matching this schema:
    {Resume.model_json_schema()}

    Important rules:
    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Parse the following resume:\n\n{resume_text}"}
    ]
    
    response = client.chat.completions.create(
        model=model_name, 
        messages=messages, 
        response_format={"type": "json_object"}
    )
    
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    return Resume(**data)

def score_canditate(job: jobDescription, resume: Resume)-> Match_result:
    # sees the job description with canditate resume and output a score

    prompt = f"""
    You are an HR recruiter. Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    
    Return JSON matching this schema:
    {Match_result.model_json_schema()}

    Give me:
    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """

    messages = [
        {'role': 'user', 'content': prompt}
    ]

    response = client.chat.completions.create(
        model = model_name, messages = messages, response_format = {"type": "json_object"}
    )
    raw_data = response.choices[0].message.content
    match_result = json.loads(raw_data)
    return Match_result(**match_result)

### Main code 

def evaluate_candidates_api(job_description_text: str, resume_files: list[tuple[str, bytes]]) -> dict:
    parsed_job = parse_job_description(job_description_text)
    
    all_results = []
    
    for filename, file_bytes in resume_files:
        import io
        file_obj = io.BytesIO(file_bytes)
        
        resume_text = read_resume(file_obj, filename)
        if not resume_text:
            continue
        
        parsed_resume = parse_resume(resume_text)
        # Avoid rate limits
        time.sleep(2)
        
        matched_content = score_canditate(parsed_job, parsed_resume)
        # Avoid rate limits
        time.sleep(2)
        
        all_results.append({
            'Name': parsed_resume.name or "Unknown Candidate",
            'Score': matched_content.score,
            'details': matched_content.details
        })
        
    if all_results:
        all_results.sort(key=lambda c: c["Score"], reverse=True)
        
        top_2 = all_results[:2]
        worst_2 = all_results[-2:] if len(all_results) > 2 else []
        
        return {
            "success": True,
            "target_role": parsed_job.role,
            "minimum_experience": parsed_job.minimum_experience,
            "top_candidates": top_2,
            "lowest_candidates": worst_2,
            "all_results": all_results
        }
    
    return {"success": False, "message": "No resumes successfully processed."}


