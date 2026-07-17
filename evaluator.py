import os 
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document
import json
import time

load_dotenv()
API_KEY = os.getenv("GROQ_API_KEY")

if not API_KEY:
    raise ValueError("Api key not found!!")

client = Groq(api_key = API_KEY)
model_name = 'llama-3.3-70b-versatile'

### Schemas (blueprints)
class jobDescription(BaseModel):
    role:'str'
    required_skills : list[str]
    preffered_skills : list[str]
    minimum_experience : float | None
    educational_requirement: list[str]
    responsibilities: list[str]

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str| None
    email: str | None
    phone: str | None
    total_experience_year: float | None
    skills : list[str] = []
    experiences: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

class Match_result(BaseModel):
    score: float 
    details: dict

### Reading the resumes 
def read_resume(FilePath: Path)-> str| None:
    ## function have the file location, sees if it is pdf or word and read it accordingly and return the contains
    ext = FilePath.suffix.lower()
    if ext == ".pdf":
        return read_pdf(FilePath)
    elif ext == ".docx":
        return read_docx(FilePath)
    else:
        return None

def read_pdf(FilePath: Path) -> str:
    ## read the resume of pdf format and return the content in string format
    text = ""
    content = PdfReader(FilePath)
    for page in content.pages:
        if page:
            extracted = page.extract_text()
            if extracted:
                text += extracted + '\n'
    return text

def read_docx(filePath: Path) -> str:
    ## read the resume of docx format and return the content in string format
    text = ""
    content = Document(filePath)
    for paragraph in content.paragraphs:
        if paragraph:
            text += paragraph.text + '\n'
    
    for table in content.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + '\n'
    return text

### Functions to interact with LLM 

def parse_job_description(job_text : str)-> jobDescription:
    ## Extracts structured json format file from specified raw job description
    system_prompt = f'''You are an expert HR assistant.
    Your job is to analyze job descriptions and extract structured information from them.
    Return ONLY valid JSON matching this schema:
    {jobDescription.model_json_schema()}
    
    IMPORTANT:
    Do NOT return the schema itself. Do NOT return fields like "properties", "title" or "type".
    Fill the schema with actual information extracted from the job description.
    If minimum experience is not mentioned, return null.
    If information for a list is missing, return an empty list.
    Do not invent information.'''

    messages = [
        { "role": 'system', 'content': system_prompt},
        {'role': "user", 'content': f"analyze the following job descriptions and do the steps you are told:\n\n{job_text}"}
    ]

    response = client.chat.completions.create(model = model_name, messages = messages, response_format = {"type": "json_object"})
    raw_json = response.choices[0].message.content
    job_data = json.loads(raw_json)
    return  jobDescription(**job_data)


def parse_resume(resume_text: str) -> Resume:
    """Extracts structured JSON data from a raw Resume."""
    system_prompt = f"""
    You are an expert resume parser.
    Extract information from the resume based on its meaning, not only based on exact section headings.
    Return ONLY valid JSON matching this schema:
    {Resume.model_json_schema()}

    Important rules:
    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Parse the following resume:\n\n{resume_text}"}
    ]
    
    response = client.chat.completions.create(
        model=model_name, 
        messages=messages, 
        response_format={"type": "json_object"}
    )
    
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    return Resume(**data)

def score_canditate(job: jobDescription, resume: Resume)-> Match_result:
    # sees the job description with canditate resume and output a score

    prompt = f"""
    You are an HR recruiter. Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    
    Return JSON matching this schema:
    {Match_result.model_json_schema()}

    Give me:
    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """

    messages = [
        {'role': 'user', 'content': prompt}
    ]

    response = client.chat.completions.create(
        model = model_name, messages = messages, response_format = {"type": "json_object"}
    )
    raw_data = response.choices[0].message.content
    match_result = json.loads(raw_data)
    return Match_result(**match_result)

### Main code 

raw_job_description = """
Description
    A career in IBM Consulting is built on long-term client relationships and close collaboration worldwide. You’ll work with leading companies across industries, helping them shape their hybrid cloud and AI journeys. With support from our strategic partners, robust IBM technology, and Red Hat, you’ll have the tools to drive meaningful change and accelerate client impact. At IBM Consulting, curiosity fuels success. You’ll be encouraged to challenge the norm, explore new ideas, and create innovative solutions that deliver real results. Our culture of growth and empathy focuses on your long-term career development while valuing your unique skills and experiences.

Your Role And Responsibilities

As a Data Engineer with expertise in Machine Learning, you will apply Machine Learning concepts and techniques to address business challenges. You will leverage your skills to drive informed decision-making in the organization. Your primary responsibilities will include:

Develop Machine Learning Solutions: Apply Machine Learning concepts and techniques to address business challenges, interpreting statistical data and identifying relevant features to inform solution development.
Evaluate Algorithm Performance: Choose appropriate algorithms and evaluate their performance using relevant metrics, ensuring that solutions meet business needs and drive informed decision-making.
Communicate Results: Clearly communicate the results of Machine Learning initiatives to stakeholders, providing actionable insights that inform business decisions.
Implement Machine Learning Techniques: Collaborate with stakeholders to implement Machine Learning techniques that drive business value, selecting and applying relevant methodologies to achieve desired outcomes.

Preferred Education

Master's Degree

Required Technical And Professional Expertise

Exposure to Machine Learning Concepts: Familiarity with applying Machine Learning concepts and techniques to address business challenges, including interpreting statistical data and identifying relevant features.
Algorithm Development Experience: Experience working with algorithms, including choosing appropriate algorithms and evaluating their performance using relevant metrics.
Data Analysis Skills: Ability to interpret statistical data and identify relevant features to inform solution development.
Machine Learning Implementation: Exposure to implementing Machine Learning techniques, including collaborating with stakeholders to select and apply relevant methodologies.
Technical Communication: Experience communicating complex technical results to stakeholders, providing actionable insights that inform business decisions.

Preferred Technical And Professional Experience

Advanced Algorithm Development: Experience working with complex algorithms, including evaluating their performance using relevant metrics and fine-tuning for optimal results.
Data Visualization Techniques: Exposure to data visualization tools and techniques, enabling effective communication of Machine Learning results to stakeholders.
Specialized Machine Learning Tools: Familiarity with specialized Machine Learning tools and technologies, such as those used for natural language processing or computer vision.
"""
print("Parsing raw job description to the LLM")
parsed_job = parse_job_description(raw_job_description)
print(f"Target role = {parsed_job.role}")
print(f"Minimum experience needed: {parsed_job.minimum_experience}")
print('------------------------------------------------------------')

resume_folder = Path('.resumes')
all_results = []

if not resume_folder.exists():
    print(f"Directory {resume_folder} is not present, please make a folder and add the resumes ")

else:
    for file_path in resume_folder.iterdir():
        if file_path.suffix.lower() not in [".pdf", ".docx"]:
            continue
            
        print(f"Processing: {file_path.name} ")
        resume_text = read_resume(file_path)
        if not resume_text:
            print("Failed to read resume")
            continue
        
        parsed_resume= parse_resume(resume_text)

        time.sleep(5)

        matched_content = score_canditate(parsed_job, parsed_resume)

        time.sleep(5)

        print(f"Score for {parsed_resume.name} : {matched_content.score}")
        all_results.append({
            'Name': parsed_resume.name or "Unknown Candidate",
            'Score': matched_content.score,
            "details": matched_content.details
        })

if all_results:
        all_results.sort(key=lambda c: c["Score"], reverse=True)
        
        top_2 = all_results[:2]
        worst_2 = all_results[-2:] if len(all_results) > 2 else []

        print("\n================ TOP 2 CANDIDATES ================")
        for candidate in top_2:
            print(f"{candidate['Name']} - {candidate['Score']}%")
            print(json.dumps(candidate["details"], indent=2))

        if worst_2:
            print("\n================ LOWEST 2 CANDIDATES ================")
            for candidate in worst_2:
                print(f"{candidate['Name']} - {candidate['Score']}%")
                print(json.dumps(candidate["details"], indent=2))

